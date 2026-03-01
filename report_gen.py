from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_rtl(paragraph):
    """ضبط اتجاه الفقرة من اليمين إلى اليسار (RTL)."""
    p_pr = paragraph._element.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    bidi.set(qn('w:val'), '1')
    p_pr.append(bidi)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

def make_arabic_report():
    doc = Document()
    
    # تحويل الصفحة بالكامل لجهة اليمين
    for section in doc.sections:
        section.header_distance = Inches(0.5)
        section.footer_distance = Inches(0.5)

    # العنوان الرئيسي
    heading = doc.add_heading("تقرير مشروع أداة الربط المكاني والوصفي GIS", 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # معلومات المشروع
    p = doc.add_paragraph(f"رابط GitHub: https://github.com/mohamad1991277-spec/gis-join-tool")
    set_rtl(p)
    p = doc.add_paragraph(f"رابط التطبيق المباشر: https://gis-join-tool.streamlit.app")
    set_rtl(p)
    p = doc.add_paragraph("إعداد الطالب: أنس محمد زقوت")
    set_rtl(p)
    p = doc.add_paragraph("التاريخ: مارس 2026")
    set_rtl(p)

    # 1. فكرة التطبيق
    h1 = doc.add_heading("1. فكرة التطبيق", level=1)
    set_rtl(h1)
    p1 = doc.add_paragraph(
        "هذا التطبيق هو أداة ويب متخصصة في نظم المعلومات الجغرافية (GIS) مصممة لإجراء عمليتين حيويتين في تحليل البيانات المكانية: "
        "الربط المكاني (Spatial Join) والربط الوصفي (Attribute Join). تم بناء الأداة باستخدام لغة Python وإطار عمل Streamlit، "
        "وتهدف إلى توفير واجهة سهلة الاستخدام للباحثين والطلاب لدمج الطبقات الجغرافية دون الحاجة إلى برامج مكتبية ثقيلة مثل ArcGIS أو QGIS."
    )
    set_rtl(p1)
    
    p_feat = doc.add_paragraph("الميزات الرئيسية:")
    set_rtl(p_feat)
    features = [
        "رفع الملفات بشكل مستقل لصيغ Shapefile (ZIP) و GeoJSON.",
        "معاينة تفاعلية للخرائط لكلا الطبقتين المدخلتين.",
        "عرض جداول البيانات لاطلاع سريع على الخصائص.",
        "ربط مكاني متقدم يدعم 8 علاقات هندسية (كالتقاطع، الاحتواء، إلخ).",
        "ربط وصفي قياسي يدعم استراتيجيات دمج متعددة (Left, Inner, Right, Outer).",
        "تنزيل مباشر لنتيجة الربط بصيغة GeoJSON العالمية."
    ]
    for feat in features:
        p = doc.add_paragraph(f"• {feat}")
        set_rtl(p)

    # 2. خطوات التنفيذ
    h2 = doc.add_heading("2. خطوات التنفيذ", level=1)
    set_rtl(h2)
    
    h2_1 = doc.add_heading("2.1 نظام معالجة الملفات", level=2)
    set_rtl(h2_1)
    p2_1 = doc.add_paragraph(
        "يعتمد النظام على آلية تحميل مزدوجة. بالنسبة لملفات Shapefile، ونظراً لكونها تتكون من ملفات متعددة، "
        "يقوم التطبيق باستخراج ملف ZIP المرفوع إلى مجلد مؤقت باستخدام مكتبة tempfile ثم يحدد ملف .shp الرئيسي "
        "ويقرأه كـ GeoDataFrame. أما ملفات GeoJSON فيتم قراءتها مباشرة من الذاكرة لضمان أقصى كفاءة."
    )
    set_rtl(p2_1)

    h2_2 = doc.add_heading("2.2 الربط المكاني الاستراتيجي", level=2)
    set_rtl(h2_2)
    p2_2 = doc.add_paragraph(
        "يضمن تنفيذ الربط المكاني سلامة البيانات من خلال المزامنة التلقائية لأنظمة الإحداثيات (CRS). "
        "إذا كانت الطبقات تستخدم مساقط مختلفة، يتم إعادة إسقاط الطبقة الثانوية لتطابق الطبقة الأساسية قبل البدء في العمليات الهندسية."
    )
    set_rtl(p2_2)

    h2_3 = doc.add_heading("2.3 واجهة المستخدم ورسم الخرائط", level=2)
    set_rtl(h2_3)
    p2_3 = doc.add_paragraph(
        "تم تصميم الواجهة باستخدام CSS مخصص يتناسب مع تصورات الـ GIS. "
        "تعتمد الخرائط التفاعلية على مكتبة Folium، مدمجة في Streamlit عبر streamlit-folium. "
        "تضبط كل خريطة حدودها تلقائياً لتناسب نطاق البيانات المعروضة."
    )
    set_rtl(p2_3)

    # 4. المشاكل والتحديات
    h4 = doc.add_heading("4. التحديات التي واجهتني أثناء التطوير", level=1)
    set_rtl(h4)
    
    challenges = [
        "المشكلة 1: التعامل مع ملفات Shapefile المتعددة في بيئة خادم ويب. الحل: استخدام مكتبات zipfile و tempfile لإنشاء مجلدات مؤقتة تُحذف تلقائياً بعد المعالجة.",
        "المشكلة 2: إدارة حالة الجلسة لنتائج الـ GIS. الحل: تخزين نتائج الـ GeoDataFrames في st.session_state للحفاظ على استمرارية البيانات.",
        "المشكلة 3: دعم الهندسات المتعددة في الرسم. الحل: تطوير دالة render_map قوية تعالج الاستثناءات الهندسية مع الحفاظ على وظيفة التلميحات التفاعلية."
    ]
    for ch in challenges:
        p = doc.add_paragraph(ch)
        set_rtl(p)

    # 5. المظهر الاحترافي
    h5 = doc.add_heading("5. المظهر الاحترافي", level=1)
    set_rtl(h5)
    p5 = doc.add_paragraph(
        "تم إخفاء القوائم الافتراضية لأطر العمل (Streamlit Menu) والأزرار البرمجية من الواجهة باستخدام أكواد CSS مخصصة، "
        "وذلك لضمان ظهور التطبيق كموقع ويب مستقل ومحتوي على تصميم خاص وفريد، مما يعزز من المظهر الاحترافي للمشروع."
    )
    set_rtl(p5)

    # الخاتمة
    h6 = doc.add_heading("خلاصة", level=1)
    set_rtl(h6)
    conc = doc.add_paragraph(
        "نجح المشروع في تلبية جميع المتطلبات الوظيفية، مقدماً أداة GIS احترافية. "
        "يسد التطبيق الفجوة بين برامج الـ GIS المعقدة وواجهات الويب البسيطة، مما يجعل التحليل المكاني أكثر سهولة وكفاءة."
    )
    set_rtl(conc)

    doc.save("GIS_Project_Report_Arabic.docx")
    print("Report saved as GIS_Project_Report_Arabic.docx")

if __name__ == "__main__":
    make_arabic_report()
